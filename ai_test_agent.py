"""
=============================================================================
AI TEST AGENT v2.0 - SCHEDULED WITH EMAIL, REPORTING & JIRA INTEGRATION
Intelligent Automated Testing Agent for Automation Anywhere Website
Created by: NVISH Solutions
Date: March 2026
=============================================================================
"""

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import Select
import time
import os
import json
import smtplib
import schedule
import requests
from requests.auth import HTTPBasicAuth
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email import encoders
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import traceback
import streamlit as st
from dotenv import load_dotenv
import os

load_dotenv()  # loads .env file

# Works for both local .env and Streamlit Secrets
def get_config(key, default=None):
    try:
        return st.secrets[key]
    except:
        return os.getenv(key, default)

# =============================================================================
# EMAIL CONFIGURATION - OUTLOOK SETTINGS
# =============================================================================
EMAIL_CONFIG = {
    "sender_email": get_config("SMTP_EMAIL"),
    "sender_password": get_config("SMTP_PASSWORD"),
    "recipients": get_config("SMTP_RECIPIENTS", "").split(","),
    "smtp_server": get_config("SMTP_SERVER"),
    "smtp_port": int(get_config("SMTP_PORT", 587))
}


# =============================================================================
# JIRA CONFIGURATION
# =============================================================================
JIRA_CONFIG = {
    "enabled": get_config("JIRA_ENABLED", "True") == "True",
    "base_url": get_config("JIRA_BASE_URL"),
    "email": get_config("JIRA_EMAIL"),
    "api_token": get_config("JIRA_API_TOKEN"),
    "project_key": get_config("JIRA_PROJECT_KEY"),
    "issue_type": get_config("JIRA_ISSUE_TYPE", "Task"),
    "assignee_email": get_config("JIRA_ASSIGNEE_EMAIL"),
    "labels": get_config("JIRA_LABELS", "").split(","),
    "custom_fields": {
        "digital_marketing_sla": get_config("JIRA_SLA"),
        "product_backlog_ready": get_config("JIRA_BACKLOG_READY"),
    }
}

# =============================================================================
# SCHEDULE CONFIGURATION
# -----------------------------------------------------------------------------
# The `schedule` library uses YOUR SYSTEM CLOCK (not UTC).
#   - If your Windows clock shows IST time  →  keep "19:00"
#   - If your Windows clock shows UTC time  →  change to "13:30" (UTC+5:30)
# To verify: open CMD and run `time` — it shows your system clock.
# =============================================================================
SCHEDULE_TIME = "19:00"   # 7:00 PM IST — change to "13:30" if clock is UTC


class AITestAgentScheduled:
    def __init__(self, agent_name="QA-Agent", run_headless=True):
        self.run_headless = run_headless
        self.agent_name = agent_name
        self.version = "2.0.0"
        self.start_time = datetime.now()
        self.knowledge_base = self.load_knowledge_base()
        self.test_results = []
        self.decisions_made = []
        self.learnings = []
        self.screenshots = []
        self.issues_found = []

        self.desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        self.reports_folder = os.path.join(self.desktop, "AI_Agent_Reports")
        self.screenshots_folder = os.path.join(self.reports_folder, "screenshots")

        if not os.path.exists(self.reports_folder):
            os.makedirs(self.reports_folder)
        if not os.path.exists(self.screenshots_folder):
            os.makedirs(self.screenshots_folder)

        print("\n" + "=" * 70)
        print("AI TEST AGENT v2.0 - SCHEDULED EXECUTION")
        print("=" * 70)
        print(f"   Agent Name    : {self.agent_name}")
        print(f"   Version       : {self.version}")
        print(f"   Started At    : {self.start_time.strftime('%Y-%m-%d %H:%M:%S')}")
        print("=" * 70)

    def load_knowledge_base(self):
        kb_path = os.path.join(os.path.expanduser("~"), "Desktop", "agent_knowledge.json")
        if os.path.exists(kb_path):
            try:
                with open(kb_path, 'r') as f:
                    return json.load(f)
            except:
                pass
        return {"element_patterns": {}, "successful_strategies": {}, "failure_patterns": {}}

    def save_knowledge_base(self):
        kb_path = os.path.join(os.path.expanduser("~"), "Desktop", "agent_knowledge.json")
        try:
            with open(kb_path, 'w') as f:
                json.dump(self.knowledge_base, f, indent=2)
        except:
            pass

    def log_decision(self, decision_type, context, action_taken, result):
        self.decisions_made.append({
            "timestamp": datetime.now().strftime("%H:%M:%S"),
            "type": decision_type,
            "context": context,
            "action": action_taken,
            "result": result
        })

    def capture_screenshot(self, driver, test_name, status="failure"):
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = test_name.replace(" ", "_").replace("-", "_").replace("/", "_")
            filename = f"{status}_{safe_name}_{timestamp}.png"
            filepath = os.path.join(self.screenshots_folder, filename)

            driver.save_screenshot(filepath)
            self.screenshots.append({
                "test_name": test_name,
                "filepath": filepath,
                "filename": filename,
                "status": status,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "url": driver.current_url
            })
            print(f"      Screenshot captured: {filename}")
            return filepath
        except Exception as e:
            print(f"      Failed to capture screenshot: {str(e)[:30]}")
            return None

    def add_issue(self, test_name, issue_type, description, screenshot_path=None, url=None, additional_info=None):
        self.issues_found.append({
            "test_name": test_name,
            "issue_type": issue_type,
            "description": description,
            "screenshot_path": screenshot_path,
            "url": url,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "additional_info": additional_info or {}
        })

    def smart_find_element(self, driver, element_name, strategies):
        print(f"      Searching for: {element_name}")

        if element_name in self.knowledge_base.get("element_patterns", {}):
            known_selector = self.knowledge_base["element_patterns"][element_name]
            try:
                element = driver.find_element(By.CSS_SELECTOR, known_selector)
                if element.is_displayed():
                    print(f"      Found using learned pattern")
                    self.log_decision("ELEMENT_FIND", element_name, "Used knowledge base", "SUCCESS")
                    return element
            except:
                pass

        for i, (selector_type, selector) in enumerate(strategies):
            try:
                if selector_type == "css":
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                elif selector_type == "xpath":
                    elements = driver.find_elements(By.XPATH, selector)
                elif selector_type == "text":
                    elements = driver.find_elements(By.PARTIAL_LINK_TEXT, selector)
                elif selector_type == "id":
                    elements = driver.find_elements(By.ID, selector)
                elif selector_type == "name":
                    elements = driver.find_elements(By.NAME, selector)
                else:
                    continue

                for element in elements:
                    if element.is_displayed():
                        if selector_type == "css":
                            self.knowledge_base.setdefault("element_patterns", {})[element_name] = selector
                            self.learnings.append(f"Learned: '{element_name}' found with '{selector}'")

                        print(f"      Found using strategy {i+1}: {selector_type}")
                        self.log_decision("ELEMENT_FIND", element_name, f"Strategy {i+1}: {selector}", "SUCCESS")
                        return element
            except:
                continue

        print(f"      Element not found after trying {len(strategies)} strategies")
        self.log_decision("ELEMENT_FIND", element_name, "All strategies failed", "FAILED")
        return None

    def handle_popups(self, driver):
        popups_handled = []

        try:
            cookie_btn = driver.find_element(By.ID, "onetrust-accept-btn-handler")
            if cookie_btn.is_displayed():
                cookie_btn.click()
                popups_handled.append("cookie")
                time.sleep(1)
        except:
            pass

        try:
            close_btns = driver.find_elements(By.CSS_SELECTOR, "[class*='close'], [aria-label*='close'], [aria-label*='Close']")
            for btn in close_btns[:3]:
                try:
                    if btn.is_displayed() and btn.is_enabled():
                        btn.click()
                        popups_handled.append("modal")
                        time.sleep(0.5)
                except:
                    pass
        except:
            pass

        try:
            ActionChains(driver).send_keys(Keys.ESCAPE).perform()
        except:
            pass

        return popups_handled

    def add_result(self, test_name, status, message, details=None):
        self.test_results.append({
            "test_name": test_name,
            "status": status,
            "message": message,
            "details": details or {},
            "timestamp": datetime.now().strftime("%H:%M:%S")
        })

    def fill_form_field(self, driver, field_name, value, strategies):
        element = self.smart_find_element(driver, field_name, strategies)
        if element:
            try:
                element.clear()
                element.send_keys(value)
                return True
            except:
                pass
        return False

    def select_dropdown(self, driver, field_name, strategies, value_to_select=None):
        element = self.smart_find_element(driver, field_name, strategies)
        if element:
            try:
                select = Select(element)
                if value_to_select:
                    try:
                        select.select_by_visible_text(value_to_select)
                        return True
                    except:
                        pass
                for option in select.options:
                    if option.get_attribute("value") and option.get_attribute("value") != "":
                        select.select_by_value(option.get_attribute("value"))
                        return True
            except:
                try:
                    element.click()
                    time.sleep(0.5)
                    options = driver.find_elements(By.CSS_SELECTOR, "[role='option'], li[class*='option'], div[class*='option']")
                    if options:
                        options[0].click()
                        return True
                except:
                    pass
        return False

    def run_test_suite(self, mode: str = "all"):

        print("\n" + "=" * 70)
        print("STARTING INTELLIGENT TEST EXECUTION")
        print("=" * 70)

        mode = (mode or "all").lower()
        run_navigation  = mode in ("all", "navigation")
        run_form        = mode in ("all", "form")
        run_performance = mode in ("all", "performance")
        
        run_mobile      = mode in ("all", "mobile")
        
        
        
        import platform
        options = webdriver.ChromeOptions()

        # Use headless based on parameter or if running on Linux cloud
        if self.run_headless or platform.system() == "Linux":
            options.add_argument("--headless=new")
            print("   INFO - Running in HEADLESS mode")
        else:
            print("   INFO - Running in HEADED mode (browser visible)")

        options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
        options.add_argument("--start-maximized")
        options.add_argument("--disable-notifications")
        options.add_argument("--disable-popup-blocking")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        options.add_argument("--disable-extensions")
        options.add_argument("--window-size=1920,1080")
        options.add_argument("--force-device-scale-factor=1")
        options.add_experimental_option('excludeSwitches', ['enable-logging', 'enable-automation'])
        options.add_experimental_option('useAutomationExtension', False)

        if platform.system() == "Linux":
            if os.path.exists("/usr/bin/google-chrome"):
                options.binary_location = "/usr/bin/google-chrome"
            elif os.path.exists("/usr/bin/chromium-browser"):
                options.binary_location = "/usr/bin/chromium-browser"
            service = Service("/usr/bin/chromedriver")
        else:
            service = Service(ChromeDriverManager().install())

        driver = webdriver.Chrome(service=service, options=options)
        driver.implicitly_wait(10)
        wait = WebDriverWait(driver, 30)
        try:
            # =========================
            # NAVIGATION GROUP
            # (TEST 1, 2, 3, 4, 5)
            # =========================
            if run_navigation:
                # TEST 1: Website Accessibility
                print("\n" + "-" * 70)
                print("TEST 1: Website Accessibility")
                print("-" * 70)

                driver.get("https://www.automationanywhere.com/")
                time.sleep(4)
                self.handle_popups(driver)

                if "Automation" in driver.title:
                    print("   PASS - Website loaded successfully")
                    self.add_result("Website Accessibility", "PASS", driver.title)
                else:
                    print("   FAIL - Website not loaded")
                    screenshot = self.capture_screenshot(driver, "Website_Accessibility", "failure")
                    self.add_result("Website Accessibility", "FAIL", "Title mismatch")
                    self.add_issue(
                        "Website Accessibility",
                        "Page Load Failure",
                        f"Expected 'Automation' in title, got: {driver.title}",
                        screenshot,
                        driver.current_url
                    )

                # TEST 2: Navigation Links
                print("\n" + "-" * 70)
                print("TEST 2: Navigation Links")
                print("-" * 70)

                nav_items = [
                    {"name": "Products",  "strategies": [("text", "Products"),  ("css", "a[href*='product']")]},
                    {"name": "Solutions", "strategies": [("text", "Solutions"), ("css", "a[href*='solution']")]},
                    {"name": "Resources", "strategies": [("text", "Resources"), ("text", "Learn")]},
                    {"name": "Company",   "strategies": [("text", "Company"),   ("text", "About")]}
                ]

                for item in nav_items:
                    element = self.smart_find_element(driver, f"nav_{item['name']}", item['strategies'])
                    if element:
                        print(f"   PASS - '{item['name']}' found")
                        self.add_result(f"Nav - {item['name']}", "PASS", "Found")
                    else:
                        try:
                            header = driver.find_element(By.CSS_SELECTOR, "header, nav")
                            if item['name'].lower() in header.get_attribute("innerHTML").lower():
                                print(f"   PASS - '{item['name']}' found in header")
                                self.add_result(f"Nav - {item['name']}", "PASS", "Found in HTML")
                            else:
                                print(f"   FAIL - '{item['name']}' not found")
                                self.add_result(f"Nav - {item['name']}", "FAIL", "Not found")
                        except:
                            self.add_result(f"Nav - {item['name']}", "FAIL", "Not found")

                
                                # TEST 3: Search Functionality
                print("\n" + "-" * 70)
                print("TEST 3: Search Functionality")
                print("-" * 70)

                search_success = False

                try:
                    # Step 1: Navigate directly to the search page
                    print("   Step 1: Navigating to search page...")
                    driver.get("https://www.automationanywhere.com/search")
                    time.sleep(3)
                    self.handle_popups(driver)
                    
                    # Step 2: Find the search input
                    print("   Step 2: Finding search input...")
                    search_input = None
                    
                    search_selectors = [
                        'input[type="search"]',
                        'input[name="s"]',
                        'input[placeholder*="Search"]',
                        'input[class*="search"]',
                    ]
                    
                    for selector in search_selectors:
                        try:
                            search_input = WebDriverWait(driver, 5).until(
                                EC.visibility_of_element_located((By.CSS_SELECTOR, selector))
                            )
                            if search_input:
                                print(f"   Found search input with: {selector}")
                                break
                        except:
                            continue
                    
                    if not search_input:
                        raise Exception("Search input not found on search page")
                    
                    # Step 3: Enter search query
                    print("   Step 3: Entering search query...")
                    query = "rpa"
                    search_input.clear()
                    search_input.send_keys(query)
                    print(f"   Entered search query: '{query}'")
                    time.sleep(1)
                    
                    # Step 4: Submit the search
                    print("   Step 4: Submitting search...")
                    
                    submitted = False
                    
                    # Try clicking submit button
                    try:
                        submit_btn = driver.find_element(By.CSS_SELECTOR, 'form button[type="submit"]')
                        driver.execute_script("arguments[0].click();", submit_btn)
                        submitted = True
                        print("   Clicked submit button")
                    except:
                        pass
                    
                    # Fallback: Press Enter
                    if not submitted:
                        try:
                            search_input.send_keys(Keys.ENTER)
                            submitted = True
                            print("   Pressed Enter key")
                        except:
                            pass
                    
                    # Fallback: JavaScript submit
                    if not submitted:
                        driver.execute_script("""
                            var form = document.querySelector('form');
                            if (form) form.submit();
                        """)
                        print("   Submitted via JavaScript")
                    
                    time.sleep(4)
                    
                    # Step 5: Verify search was performed
                    print("   Step 5: Verifying search results...")
                    current_url = driver.current_url
                    print(f"   Current URL: {current_url}")
                    
                    # Check URL contains search query
                    if "search" in current_url.lower() and ("s=" in current_url or "q=" in current_url or "query" in current_url):
                        print(f"   PASS - Search performed successfully!")
                        
                        # Additional check: Verify results are displayed
                        results_found = driver.execute_script("""
                            var pageText = document.body.innerText.toLowerCase();
                            return pageText.includes('result') || 
                                   document.querySelectorAll('[class*="result"], article, [class*="card"]').length > 0;
                        """)
                        
                        if results_found:
                            print("   Search results are displayed on the page")
                        
                        self.add_result(
                            "Search Functionality",
                            "PASS",
                            f"Search performed successfully with query '{query}'",
                            {"url": current_url, "results_displayed": results_found},
                        )
                        self.capture_screenshot(driver, "Search_Results_Page", "pass")
                        search_success = True
                    else:
                        print(f"   FAIL - Search URL not updated properly")
                        self.capture_screenshot(driver, "Search_URL_Issue", "failure")

                except Exception as e:
                    print(f"   Exception during search test: {str(e)}")
                    self.capture_screenshot(driver, "Search_Exception", "failure")

                if not search_success:
                    print("   FAIL - Search functionality test failed")
                    screenshot = self.capture_screenshot(driver, "Search_Failed", "failure")
                    self.add_result(
                        "Search Functionality",
                        "FAIL",
                        "Search functionality test failed",
                        {"screenshot": screenshot, "url": driver.current_url},
                    )
                    self.add_issue(
                        "Search Functionality",
                        "Search Test Failed",
                        "Could not perform search on the search page",
                        screenshot,
                        driver.current_url,
                    )

                # Return to homepage for next test
                driver.get("https://www.automationanywhere.com/")
                time.sleep(3)
                self.handle_popups(driver)

                
                # TEST 4: Products Page


                # TEST 4: Products Page
                print("\n" + "-" * 70)
                print("TEST 4: Products Page")
                print("-" * 70)


                driver.get("https://www.automationanywhere.com/products")
                time.sleep(4)
                self.handle_popups(driver)

                if "product" in driver.current_url.lower():
                    print("   PASS - Products page loaded")
                    self.add_result("Products Page", "PASS", driver.current_url)
                else:
                    print("   FAIL - Products page not loaded")
                    self.add_result("Products Page", "FAIL", "URL mismatch")

                # TEST 5: Solutions Page
                print("\n" + "-" * 70)
                print("TEST 5: Solutions Page")
                print("-" * 70)

                driver.get("https://www.automationanywhere.com/solutions")
                time.sleep(4)
                self.handle_popups(driver)

                print("   PASS - Solutions page loaded")
                self.add_result("Solutions Page", "PASS", driver.current_url)

            # =========================
            # =========================
            # FORM GROUP
            # (TEST 6)
            # =========================
            if run_form:
                print("\n" + "-" * 70)
                print("TEST 6: Demo Request Form")
                print("-" * 70)

                driver.get("https://www.automationanywhere.com/request-live-demo")

                # Wait for page to fully load dynamically
                try:
                    wait.until(lambda d: d.execute_script("return document.readyState") == "complete")
                    print("   INFO - Page loaded")
                except:
                    print("   WARN - Page load timeout, proceeding anyway")

                self.handle_popups(driver)
                driver.execute_script("window.scrollBy(0, 400)")

                # Wait for form to appear - moves on immediately when found
                try:
                    form_wait = WebDriverWait(driver, 20)
                    form_wait.until(EC.presence_of_element_located((By.CSS_SELECTOR,
                        "form.mktoForm, input[name*='FirstName'], input[name='Email'], input[type='email']")))
                    print("   INFO - Form detected, proceeding")
                except:
                    print("   WARN - Form not detected after 20s, proceeding anyway")

                test_data = {
                    "first_name": "Test",
                    "last_name": "Automation",
                    "email": "test@nvish.com",
                    "phone": "9876543210",
                    "company": "NVISH Solutions",
                }

                fields_filled = 0

                # Helper - wait for element dynamically then fill
                def wait_and_fill(field_name, css_selector, value, result_label):
                    try:
                        element = WebDriverWait(driver, 15).until(
                            EC.presence_of_element_located((By.CSS_SELECTOR, css_selector))
                        )
                        driver.execute_script("arguments[0].scrollIntoView(true);", element)
                        element.clear()
                        element.send_keys(value)
                        print(f"   PASS - {result_label} filled")
                        self.add_result(f"Form - {result_label}", "PASS", value)
                        return True
                    except:
                        print(f"   FAIL - {result_label} field not found")
                        self.add_result(f"Form - {result_label}", "FAIL", "Field not found")
                        return False

                # Helper - wait for dropdown dynamically
                def wait_and_find_select(field_name, css_selector, timeout=15):
                    try:
                        element = WebDriverWait(driver, timeout).until(
                            EC.presence_of_element_located((By.CSS_SELECTOR, css_selector))
                        )
                        return element
                    except:
                        return None

                # ---------- 1. First Name ----------
                if wait_and_fill(
                    "first_name",
                    "input[name='FirstName'], input[name*='FirstName'], input[id*='FirstName'], input[placeholder*='First']",
                    test_data["first_name"],
                    "First Name"
                ):
                    fields_filled += 1

                # ---------- 2. Last Name ----------
                if wait_and_fill(
                    "last_name",
                    "input[name='LastName'], input[name*='LastName'], input[id*='LastName'], input[placeholder*='Last']",
                    test_data["last_name"],
                    "Last Name"
                ):
                    fields_filled += 1

                # ---------- 3. Business Email ----------
                if wait_and_fill(
                    "business_email",
                    "input[name='Email'], input[name*='Email'], input[type='email'], input[id*='Email']",
                    test_data["email"],
                    "Business Email"
                ):
                    fields_filled += 1

                # ---------- 4. Phone Number ----------
                if wait_and_fill(
                    "phone_number",
                    "input[name='Phone'], input[name*='Phone'], input[type='tel'], input[id*='Phone']",
                    test_data["phone"],
                    "Phone Number"
                ):
                    fields_filled += 1

                # ---------- 5. Company Name ----------
                if wait_and_fill(
                    "company_name",
                    "input[name='Company'], input[name*='Company'], input[id*='Company']",
                    test_data["company"],
                    "Company Name"
                ):
                    fields_filled += 1

                # ---------- 6. Number of Employees ----------
                emp_element = wait_and_find_select(
                    "number_of_employees",
                    "select[name*='Employees'], select[aria-label*='Number of Employees'], select[id*='Employees']"
                )
                if emp_element:
                    try:
                        driver.execute_script("arguments[0].scrollIntoView(true);", emp_element)
                        select_obj = Select(emp_element)
                        selected = False
                        preferred_ranges = ["100-499", "100 - 499", "100 to 499"]
                        for pref in preferred_ranges:
                            try:
                                select_obj.select_by_visible_text(pref)
                                print(f"   PASS - Number of Employees selected: '{pref}'")
                                self.add_result("Form - Number of Employees", "PASS", pref)
                                fields_filled += 1
                                selected = True
                                break
                            except:
                                continue
                        if not selected:
                            for opt in select_obj.options:
                                if opt.get_attribute("value") and opt.get_attribute("value").strip() != "":
                                    select_obj.select_by_value(opt.get_attribute("value"))
                                    print(f"   PASS - Number of Employees selected: '{opt.text}'")
                                    self.add_result("Form - Number of Employees", "PASS", opt.text)
                                    fields_filled += 1
                                    selected = True
                                    break
                        if not selected:
                            print("   WARN - Number of Employees no selectable options")
                            self.add_result("Form - Number of Employees", "WARNING", "No selectable options")
                    except Exception as e:
                        print(f"   WARN - Number of Employees failed: {str(e)[:60]}")
                        self.add_result("Form - Number of Employees", "WARNING", str(e)[:60])
                else:
                    print("   FAIL - Number of Employees dropdown not found")
                    self.add_result("Form - Number of Employees", "FAIL", "Field not found")

                # ---------- 7. Job Function ----------
                job_element = wait_and_find_select(
                    "job_function",
                    "select[name*='JobFunction'], select[name*='Job_Function'], select[aria-label*='Job Function'], select[id*='JobFunction'], select[name*='Job'], select[id*='Job']"
                )
                if job_element:
                    try:
                        driver.execute_script("arguments[0].scrollIntoView(true);", job_element)
                        selected = False
                        preferred_roles = ["Manager", "Director", "Staff", "Individual Contributor"]
                        select_obj = Select(job_element)

                        # Wait until options are populated
                        start_time = time.time()
                        while time.time() - start_time < 10:
                            options = [opt for opt in select_obj.options
                                      if opt.get_attribute("value") and opt.get_attribute("value").strip() != ""]
                            if len(options) > 0:
                                break
                            time.sleep(0.3)

                        for pref in preferred_roles:
                            try:
                                select_obj.select_by_visible_text(pref)
                                print(f"   PASS - Job Function selected: '{pref}'")
                                self.add_result("Form - Job Function", "PASS", pref)
                                fields_filled += 1
                                selected = True
                                break
                            except:
                                continue

                        if not selected and options:
                            opt = options[0]
                            select_obj.select_by_value(opt.get_attribute("value"))
                            print(f"   PASS - Job Function selected: '{opt.text}'")
                            self.add_result("Form - Job Function", "PASS", opt.text)
                            fields_filled += 1
                            selected = True

                        if not selected:
                            print("   WARN - Job Function no selectable options")
                            self.add_result("Form - Job Function", "WARNING", "No selectable options")
                    except Exception as e:
                        print(f"   WARN - Job Function failed: {str(e)[:60]}")
                        self.add_result("Form - Job Function", "WARNING", str(e)[:60])
                else:
                    print("   FAIL - Job Function dropdown not found")
                    self.add_result("Form - Job Function", "FAIL", "Field not found")

                # ---------- 8. Country/Region ----------
                country_element = wait_and_find_select(
                    "country_region",
                    "select[name*='Country'], select[aria-label*='Country/Region'], select[id*='Country']"
                )
                if country_element:
                    try:
                        driver.execute_script("arguments[0].scrollIntoView(true);", country_element)
                        select_obj = Select(country_element)
                        selected = False
                        try:
                            select_obj.select_by_visible_text("Australia")
                            print("   PASS - Country selected: 'Australia'")
                            self.add_result("Form - Country/Region", "PASS", "Australia")
                            fields_filled += 1
                            selected = True
                        except:
                            for opt in select_obj.options:
                                if "australia" in opt.text.lower():
                                    select_obj.select_by_visible_text(opt.text)
                                    print(f"   PASS - Country selected: '{opt.text}'")
                                    self.add_result("Form - Country/Region", "PASS", opt.text)
                                    fields_filled += 1
                                    selected = True
                                    break
                        if not selected:
                            print("   FAIL - Australia not found in Country dropdown")
                            self.add_result("Form - Country/Region", "FAIL", "Australia not found")
                    except Exception as e:
                        print(f"   WARN - Country dropdown failed: {str(e)[:60]}")
                        self.add_result("Form - Country/Region", "WARNING", str(e)[:60])
                else:
                    print("   FAIL - Country dropdown not found")
                    self.add_result("Form - Country/Region", "FAIL", "Field not found")

                # ---------- Submit ----------
                self.capture_screenshot(driver, "Form_Before_First_Submit", "info")

                submit_strategies = [
                    ("css", "button[type='submit']"),
                    ("css", "input[type='submit']"),
                    ("xpath", "//button[contains(text(),'Request Live Demo') or contains(text(),'Request live demo') or contains(text(),'Request demo')]"),
                ]
                submit_btn = self.smart_find_element(driver, "request_live_demo_button", submit_strategies)

                if submit_btn:
                    print("   Submit button found - clicking...")
                    self.add_result("Form - Submit Button", "PASS", "Found")
                    try:
                        driver.execute_script("arguments[0].scrollIntoView(true);", submit_btn)
                        driver.execute_script("arguments[0].click();", submit_btn)

                        # Wait for page to change dynamically instead of fixed sleep
                        try:
                            wait.until(lambda d: d.current_url != "https://www.automationanywhere.com/request-live-demo" 
                                      or "thank" in d.current_url.lower()
                                      or "thank you" in d.page_source.lower())
                            print("   INFO - Page changed after submit")
                        except:
                            print("   WARN - Page did not change after submit")
                    except Exception as e:
                        print(f"   FAIL - Submit click failed: {str(e)[:60]}")
                        self.add_result("Form - Submit Button", "FAIL", f"Click failed: {str(e)[:60]}")
                else:
                    print("   FAIL - Submit button not found")
                    self.add_result("Form - Submit Button", "FAIL", "Not found")

                # ---------- Fix validation errors and resubmit ----------
                try:
                    error_elements = driver.find_elements(
                        By.CSS_SELECTOR, ".mktoError, .mktoInvalid, [aria-invalid='true']"
                    )
                    if error_elements:
                        print(f"   INFO - {len(error_elements)} validation errors found, fixing and resubmitting")
                        for err in error_elements:
                            try:
                                parent = err.find_element(By.XPATH, ".//ancestor::*[self::div or self::span][1]")
                                inputs = parent.find_elements(By.CSS_SELECTOR, "input, select, textarea")
                                for inp in inputs:
                                    tag = inp.tag_name.lower()
                                    if tag == "input" and inp.get_attribute("type") in ["text", "tel", "email"]:
                                        driver.execute_script("arguments[0].value='';", inp)
                                        inp.send_keys("test")
                                    elif tag == "select":
                                        sel = Select(inp)
                                        for opt in sel.options:
                                            if opt.get_attribute("value") and opt.get_attribute("value").strip() != "":
                                                sel.select_by_value(opt.get_attribute("value"))
                                                break
                            except:
                                continue
                        submit_btn2 = self.smart_find_element(driver, "request_live_demo_button_retry", submit_strategies)
                        if submit_btn2:
                            driver.execute_script("arguments[0].scrollIntoView(true);", submit_btn2)
                            driver.execute_script("arguments[0].click();", submit_btn2)
                            print("   Second submit attempted")
                            # Dynamic wait after second submit
                            try:
                                wait.until(lambda d: "thank" in d.current_url.lower() 
                                          or "thank you" in d.page_source.lower())
                            except:
                                pass
                except:
                    pass

                # ---------- Verify Thank You ----------
                current_url = driver.current_url.lower()
                page_source = driver.page_source.lower()

                if "thank" in current_url or "thank-you" in current_url or "thank_you" in current_url:
                    print(f"   PASS - Redirected to Thank You page: {driver.current_url}")
                    self.add_result("Form - Thank You Redirect", "PASS", f"Redirected to: {driver.current_url}")
                    self.capture_screenshot(driver, "Form_ThankYou_Page", "pass")
                elif "thank you" in page_source or "thank-you" in page_source or "submission" in page_source:
                    print(f"   PASS - Thank You content detected: {driver.current_url}")
                    self.add_result("Form - Thank You Redirect", "PASS", f"Thank You content at: {driver.current_url}")
                    self.capture_screenshot(driver, "Form_ThankYou_Page", "pass")
                else:
                    print(f"   FAIL - No Thank You redirect. URL: {driver.current_url}")
                    screenshot = self.capture_screenshot(driver, "Form_No_ThankYou_Redirect", "failure")
                    self.add_result("Form - Thank You Redirect", "FAIL", f"Stayed at: {driver.current_url}")
                    self.add_issue(
                        "Demo Form", "Redirect Failure",
                        "After clicking Request Live Demo, page did not redirect to Thank You page.",
                        screenshot, driver.current_url
                    )

                print(f"\n   Form Summary: {fields_filled} fields filled/selected")
                self.add_result(
                    "Demo Form",
                    "PASS" if fields_filled >= 5 else "FAIL",
                    f"{fields_filled} fields filled/selected"
                )
            
            
            # =========================
            # PERFORMANCE GROUP (TEST 7)
            # =========================
            if run_performance:
                print("\n" + "-" * 70)
                print("TEST 7: Page Performance")
                print("-" * 70)

                start = time.time()
                driver.get("https://www.automationanywhere.com/")
                wait.until(lambda d: d.execute_script("return document.readyState") == "complete")
                load_time = time.time() - start

                if load_time < 5:
                    print(f"   PASS - Page loaded in {load_time:.2f}s")
                    self.add_result("Page Performance", "PASS", f"{load_time:.2f}s")
                elif load_time < 10:
                    print(f"   WARNING - Page loaded in {load_time:.2f}s")
                    self.add_result("Page Performance", "WARNING", f"{load_time:.2f}s (slow)")
                else:
                    print(f"   FAIL - Page loaded in {load_time:.2f}s")
                    self.add_result("Page Performance", "FAIL", f"{load_time:.2f}s (very slow)")

            # =========================
            # MOBILE GROUP (TEST 8)
            # =========================
            if run_mobile:
                print("\n" + "-" * 70)
                print("TEST 8: Responsive Design")
                print("-" * 70)

                viewports = [
                    {"name": "Desktop", "width": 1920, "height": 1080},
                    {"name": "Tablet",  "width": 768,  "height": 1024},
                    {"name": "Mobile",  "width": 375,  "height": 812}
                ]

                for vp in viewports:
                    try:
                        driver.set_window_size(vp['width'], vp['height'])
                        time.sleep(1)
                        driver.get("https://www.automationanywhere.com/")
                        time.sleep(2)

                        if driver.title:
                            print(f"   PASS - {vp['name']} ({vp['width']}x{vp['height']})")
                            self.add_result(f"Responsive - {vp['name']}", "PASS", f"{vp['width']}x{vp['height']}")
                        else:
                            print(f"   FAIL - {vp['name']}")
                            self.add_result(f"Responsive - {vp['name']}", "FAIL", "Page not rendered")
                    except Exception as e:
                        print(f"   WARNING - {vp['name']} skipped: {str(e)[:30]}")
                        self.add_result(f"Responsive - {vp['name']}", "WARNING", "Skipped")

                try:
                    driver.maximize_window()
                except:
                    pass

        except Exception as e:
            print(f"\nCRITICAL ERROR: {str(e)}")
            try:
                screenshot = self.capture_screenshot(driver, "Critical_Error", "failure")
            except:
                screenshot = None
            self.add_result("Critical Error", "FAIL", str(e)[:100])
            self.add_issue("Critical Error", "Unexpected Exception", str(e), screenshot, "N/A")

        finally:
            driver.quit()

        self.save_knowledge_base()
        return self.generate_reports()

    def create_issue_document(self):
        if not self.issues_found:
            return None

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = os.path.join(self.reports_folder, f"issue_report_{timestamp}.docx")

        doc = Document()

        title = doc.add_heading('AI Test Agent - Issue Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f"Date: ").bold = True
        summary.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        summary.add_run(f"Agent: ").bold = True
        summary.add_run(f"{self.agent_name}\n")
        summary.add_run(f"Total Issues Found: ").bold = True
        summary.add_run(f"{len(self.issues_found)}\n")
        summary.add_run(f"Website Tested: ").bold = True
        summary.add_run("https://www.automationanywhere.com/")

        doc.add_heading('Detailed Issues', level=1)

        for i, issue in enumerate(self.issues_found, 1):
            doc.add_heading(f"Issue #{i}: {issue['test_name']}", level=2)

            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'

            rows = table.rows
            rows[0].cells[0].text = "Issue Type"
            rows[0].cells[1].text = issue['issue_type']
            rows[1].cells[0].text = "Timestamp"
            rows[1].cells[1].text = issue['timestamp']
            rows[2].cells[0].text = "URL"
            rows[2].cells[1].text = issue['url'] or "N/A"
            rows[3].cells[0].text = "Description"
            rows[3].cells[1].text = issue['description']
            rows[4].cells[0].text = "Additional Info"
            rows[4].cells[1].text = str(issue.get('additional_info', {}))

            for row in table.rows:
                row.cells[0].paragraphs[0].runs[0].bold = True

            if issue['screenshot_path'] and os.path.exists(issue['screenshot_path']):
                doc.add_paragraph()
                doc.add_paragraph("Screenshot:").bold = True
                try:
                    doc.add_picture(issue['screenshot_path'], width=Inches(6))
                except:
                    doc.add_paragraph(f"[Screenshot: {issue['screenshot_path']}]")

            doc.add_paragraph()

        doc.add_heading('Recommendations', level=1)
        recommendations = [
            "Review failed test cases and verify if issues are due to website changes or test script issues.",
            "For intermittent failures (like search button), consider adding longer wait times or retry mechanisms.",
            "Form submission failures should be investigated for validation rule changes.",
            "Performance issues may require coordination with the development team."
        ]
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run("Generated by AI Test Agent v2.0 - NVISH Solutions").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(doc_path)
        print(f"   Issue document created: {doc_path}")
        return doc_path

    def generate_reports(self):
        end_time = datetime.now()
        duration = str(end_time - self.start_time).split('.')[0]

        total    = len(self.test_results)
        passed   = sum(1 for r in self.test_results if r['status'] == 'PASS')
        failed   = sum(1 for r in self.test_results if r['status'] == 'FAIL')
        warnings = sum(1 for r in self.test_results if r['status'] == 'WARNING')
        pass_rate = (passed / total * 100) if total > 0 else 0

        print("\n" + "=" * 70)
        print("AI AGENT EXECUTION COMPLETE")
        print("=" * 70)
        print(f"\nTEST SUMMARY:")
        print(f"   Total Tests    : {total}")
        print(f"   Passed         : {passed}")
        print(f"   Failed         : {failed}")
        print(f"   Warnings       : {warnings}")
        print(f"   Pass Rate      : {pass_rate:.1f}%")
        print(f"   Duration       : {duration}")
        print(f"   Issues Found   : {len(self.issues_found)}")
        print(f"   Screenshots    : {len(self.screenshots)}")

        issue_doc_path = None
        if self.issues_found:
            print("\nCreating issue document...")
            issue_doc_path = self.create_issue_document()

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        html_file = os.path.join(self.reports_folder, f"ai_agent_report_{timestamp}.html")

        status_color = "#28a745" if pass_rate >= 80 else "#ffc107" if pass_rate >= 50 else "#dc3545"

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>AI Test Agent Report</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial; margin: 0; background: linear-gradient(135deg, #1a1a2e, #16213e); min-height: 100vh; padding: 20px; }}
        .container {{ max-width: 1000px; margin: 0 auto; }}
        .header {{ background: linear-gradient(135deg, #667eea, #764ba2); padding: 30px; border-radius: 15px; color: white; text-align: center; margin-bottom: 20px; }}
        .header h1 {{ margin: 0 0 10px 0; font-size: 28px; }}
        .card {{ background: white; border-radius: 15px; padding: 25px; margin-bottom: 20px; box-shadow: 0 10px 40px rgba(0,0,0,0.2); }}
        .card h2 {{ color: #333; margin: 0 0 20px 0; padding-bottom: 10px; border-bottom: 2px solid #eee; }}
        .summary {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 15px; }}
        .stat {{ text-align: center; padding: 20px; border-radius: 10px; color: white; }}
        .stat.total {{ background: linear-gradient(135deg, #667eea, #764ba2); }}
        .stat.passed {{ background: linear-gradient(135deg, #11998e, #38ef7d); }}
        .stat.failed {{ background: linear-gradient(135deg, #eb3349, #f45c43); }}
        .stat.warnings {{ background: linear-gradient(135deg, #f093fb, #f5576c); }}
        .stat h3 {{ margin: 0; font-size: 36px; }}
        .stat p {{ margin: 5px 0 0 0; opacity: 0.9; }}
        .rate {{ text-align: center; padding: 20px; background: {status_color}; color: white; border-radius: 10px; font-size: 24px; margin: 20px 0; }}
        table {{ width: 100%; border-collapse: collapse; }}
        th {{ background: #667eea; color: white; padding: 12px; text-align: left; }}
        td {{ padding: 12px; border-bottom: 1px solid #eee; }}
        tr:hover {{ background: #f8f9fa; }}
        .pass {{ color: #28a745; font-weight: bold; }}
        .fail {{ color: #dc3545; font-weight: bold; }}
        .warn {{ color: #ffc107; font-weight: bold; }}
        .footer {{ text-align: center; color: rgba(255,255,255,0.6); padding: 20px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>AI Test Agent Report</h1>
            <p>Automation Anywhere Website - Daily Automated Testing</p>
            <p>Agent: {self.agent_name} | Version: {self.version} | Date: {datetime.now().strftime('%Y-%m-%d')} | Duration: {duration}</p>
        </div>

        <div class="card">
            <h2>Test Summary</h2>
            <div class="summary">
                <div class="stat total"><h3>{total}</h3><p>Total</p></div>
                <div class="stat passed"><h3>{passed}</h3><p>Passed</p></div>
                <div class="stat failed"><h3>{failed}</h3><p>Failed</p></div>
                <div class="stat warnings"><h3>{warnings}</h3><p>Warnings</p></div>
            </div>
            <div class="rate">Pass Rate: {pass_rate:.1f}%</div>
        </div>

        <div class="card">
            <h2>Test Results</h2>
            <table>
                <tr><th>#</th><th>Test</th><th>Status</th><th>Details</th><th>Time</th></tr>"""

        for i, r in enumerate(self.test_results, 1):
            cls = "pass" if r['status'] == "PASS" else "fail" if r['status'] == "FAIL" else "warn"
            html_content += f"""
                <tr>
                    <td>{i}</td>
                    <td>{r['test_name']}</td>
                    <td class="{cls}">{r['status']}</td>
                    <td>{r['message']}</td>
                    <td>{r['timestamp']}</td>
                </tr>"""

        html_content += f"""
            </table>
        </div>

        <div class="card">
            <h2>Agent Intelligence</h2>
            <p><strong>Decisions Made:</strong> {len(self.decisions_made)}</p>
            <p><strong>Patterns Learned:</strong> {len(self.learnings)}</p>
            <p><strong>Screenshots Captured:</strong> {len(self.screenshots)}</p>
            <p><strong>Issues Documented:</strong> {len(self.issues_found)}</p>
        </div>

        <div class="footer">
            <p>Generated by AI Test Agent v{self.version} - NVISH Solutions</p>
            <p>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
    </div>
</body>
</html>"""

        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)

        print(f"\nReports generated:")
        print(f"   HTML Report: {html_file}")
        if issue_doc_path:
            print(f"   Issue Document: {issue_doc_path}")

        return {
            "html_report": html_file,
            "issue_document": issue_doc_path,
            "summary": {
                "total": total,
                "passed": passed,
                "failed": failed,
                "warnings": warnings,
                "pass_rate": pass_rate,
                "issues_count": len(self.issues_found)
            },
            "tests": self.test_results,
            "issues": self.issues_found
        }


# =============================================================================
# JIRA INTEGRATION FUNCTIONS
# =============================================================================
def get_active_sprint_id(project_key, auth):
    print(f"\n   Fetching active sprint for project {project_key}...")
    try:
        url = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue/search"
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }
        params = {
            "jql": f"project = {project_key} AND sprint in openSprints()",
            "maxResults": 1,
            "fields": "customfield_10016"
        }
        response = requests.get(url, headers=headers, auth=auth, params=params)
        if response.status_code == 200:
            issues = response.json().get('issues', [])
            if issues:
                sprints = issues[0]['fields'].get('customfield_10016', [])
                if sprints:
                    for sprint in sprints:
                        if sprint.get('state') == 'active':
                            sprint_id   = sprint['id']
                            sprint_name = sprint['name']
                            print(f"   Active sprint found: [{sprint_id}] {sprint_name}")
                            return sprint_id
        print("   No active sprint found.")
        return None
    except Exception as e:
        print(f"   Could not fetch active sprint: {str(e)}")
        return None


def set_sprint_on_jira_ticket(ticket_key, sprint_id, auth):
    print(f"\n   Setting sprint [{sprint_id}] on {ticket_key}...")
    try:
        url = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue/{ticket_key}"
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }
        payload = {
            "fields": {
                "customfield_10016": sprint_id
            }
        }
        response = requests.put(url, json=payload, headers=headers, auth=auth)
        if response.status_code == 204:
            verify = requests.get(
                f"{url}?fields=customfield_10016",
                headers={"Accept": "application/json"},
                auth=auth
            )
            if verify.status_code == 200:
                saved_sprints = verify.json()['fields'].get('customfield_10016')
                if saved_sprints:
                    saved_id = saved_sprints[0].get('id') if isinstance(saved_sprints, list) else saved_sprints.get('id')
                    if str(saved_id) == str(sprint_id):
                        print(f"   Sprint verified and saved on {ticket_key}.")
                        return True
                    else:
                        print(f"   Sprint mismatch after save. Expected {sprint_id}, got {saved_id}.")
                        return False
                else:
                    print(f"   Sprint field still empty after PUT.")
                    return False
            else:
                print(f"   Could not verify sprint after setting.")
                return False
        else:
            print(f"   Failed to set sprint: {response.status_code} - {response.text[:300]}")
            return False
    except Exception as e:
        print(f"   Sprint update error: {str(e)}")
        return False


def attach_files_to_jira(ticket_key, reports, auth):
    print(f"\n   Attaching files to {ticket_key}...")
    files_to_attach = []
    report_file_keys = ['html_report', 'pdf_report', 'json_report',
                        'report_path', 'screenshot_path', 'file_path']
    for key in report_file_keys:
        file_path = reports.get(key)
        if file_path and os.path.exists(file_path):
            files_to_attach.append(file_path)
    search_dirs = [
        os.getcwd(),
        os.path.expanduser("~/Desktop"),
        os.path.dirname(os.path.abspath(__file__))
    ]
    report_extensions = ['.html', '.pdf', '.json', '.txt', '.csv']
    for search_dir in search_dirs:
        if not os.path.exists(search_dir):
            continue
        for fname in os.listdir(search_dir):
            if fname.startswith("ai_agent_report_") and \
               any(fname.endswith(ext) for ext in report_extensions):
                full_path = os.path.join(search_dir, fname)
                if full_path not in files_to_attach and os.path.exists(full_path):
                    files_to_attach.append(full_path)
    files_to_attach = list(dict.fromkeys(files_to_attach))
    if not files_to_attach:
        print("   No report files found to attach.")
        return
    print(f"   Found {len(files_to_attach)} file(s) to attach:")
    for f in files_to_attach:
        print(f"      - {os.path.basename(f)}")
    url = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue/{ticket_key}/attachments"
    headers = {
        "Accept": "application/json",
        "X-Atlassian-Token": "no-check"
    }
    for file_path in files_to_attach:
        try:
            file_name = os.path.basename(file_path)
            with open(file_path, 'rb') as f:
                response = requests.post(
                    url,
                    headers=headers,
                    auth=auth,
                    files={"file": (file_name, f, "application/octet-stream")}
                )
            if response.status_code == 200:
                print(f"   Attached: {file_name}")
            else:
                print(f"   Failed to attach {file_name}: "
                      f"{response.status_code} - {response.text[:200]}")
        except Exception as e:
            print(f"   Error attaching {file_path}: {str(e)}")


def transition_jira_ticket(ticket_key, target_status, auth):
    WORKFLOW_ORDER = [
        "Backlog",
        "In Progress",
        "In Verification",
        "Ready for Production",
        "Production Live",
        "Completed"
    ]
    WORKFLOW_STEPS = [
        ("Backlog",              "WIP"),
        ("In Progress",          "In Verification"),
        ("In Verification",      "QA Completed"),
        ("Ready for Production", "Production Live"),
        ("Production Live",      "Completed"),
    ]
    today_date = datetime.now().strftime('%Y-%m-%d')
    SCREEN_FIELDS = {
        "customfield_15267": {"value": "P1 - 1 to 3 days - Assets"},
        "customfield_15268": {"value": "Yes"},
        "customfield_15274": today_date,
        "customfield_14885": {"value": "P1 - High"},
        "assignee":          {"accountId": "5ea6d2944e41b90b7bce0c0c"},
        "priority":          {"name": "High"}
    }

    def get_transitions(ticket):
        url  = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue/{ticket}/transitions"
        hdrs = {"Accept": "application/json", "Content-Type": "application/json"}
        resp = requests.get(url, headers=hdrs, auth=auth)
        if resp.status_code == 200:
            return resp.json().get('transitions', [])
        return []

    def do_transition(ticket, transition_id, has_screen=False):
        url  = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue/{ticket}/transitions"
        hdrs = {"Accept": "application/json", "Content-Type": "application/json"}
        payload = {"transition": {"id": str(transition_id)}}
        if has_screen:
            payload["fields"] = SCREEN_FIELDS
        resp = requests.post(url, json=payload, headers=hdrs, auth=auth)
        if resp.status_code not in (204, 200):
            print(f"   Transition response {resp.status_code}: {resp.text[:300]}")
        return resp.status_code in (204, 200)

    def find_transition(transitions, name):
        for t in transitions:
            if t['name'].lower() == name.lower() or \
               t['to']['name'].lower() == name.lower():
                return t['id'], t.get('hasScreen', False)
        return None, False

    def get_current_status(ticket):
        url  = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue/{ticket}?fields=status"
        hdrs = {"Accept": "application/json"}
        resp = requests.get(url, headers=hdrs, auth=auth)
        if resp.status_code == 200:
            return resp.json()['fields']['status']['name']
        return None

    print(f"\n   Transitioning {ticket_key} to '{target_status}'...")
    try:
        status_now = get_current_status(ticket_key)
        print(f"   Current status: '{status_now}'")
        if status_now and status_now.lower() == target_status.lower():
            print(f"   Already at '{target_status}', no transition needed.")
            return True

        current_idx = next(
            (i for i, s in enumerate(WORKFLOW_ORDER)
             if s.lower() == (status_now or "").lower()), None
        )
        target_idx = next(
            (i for i, s in enumerate(WORKFLOW_ORDER)
             if s.lower() == target_status.lower()), None
        )

        if current_idx is None or target_idx is None:
            transitions = get_transitions(ticket_key)
            print(f"   Cannot map '{status_now}' -> '{target_status}' in known workflow.")
            print(f"   Available: {[(t['name'], t['to']['name']) for t in transitions]}")
            return False

        for step in range(current_idx, target_idx):
            from_name       = WORKFLOW_ORDER[step]
            to_name         = WORKFLOW_ORDER[step + 1]
            _, step_tname   = WORKFLOW_STEPS[step]
            transitions     = get_transitions(ticket_key)
            tid, has_screen = find_transition(transitions, step_tname)
            if not tid:
                tid, has_screen = find_transition(transitions, to_name)
            if not tid:
                avail = [(t['name'], t['to']['name']) for t in transitions]
                print(f"   No transition found: '{from_name}' -> '{to_name}'")
                print(f"   Available (name -> to): {avail}")
                return False
            print(f"   Step: '{from_name}' -> '{to_name}' "
                  f"via '{step_tname}' (screen={has_screen})")
            if not do_transition(ticket_key, tid, has_screen):
                print(f"   Step to '{to_name}' failed.")
                return False
            time.sleep(1)

        print(f"   Ticket reached '{target_status}' successfully.")
        return True
    except Exception as e:
        print(f"   Transition error: {str(e)}")
        return False


def create_jira_ticket(reports, agent):
    if not JIRA_CONFIG.get("enabled", False):
        print("\n   Jira integration is disabled.")
        return None

    print("\n" + "=" * 70)
    print("CREATING JIRA TICKET")
    print("=" * 70)

    try:
        summary      = reports['summary']
        total        = summary['total']
        passed       = summary['passed']
        failed       = summary['failed']
        warnings     = summary['warnings']
        pass_rate    = summary['pass_rate']
        issues_count = summary['issues_count']

        if pass_rate == 100 and failed == 0 and issues_count == 0:
            status_text     = "ALL TESTS PASSED"
            should_complete = True
        elif failed > 0 or issues_count > 0:
            status_text     = f"{failed} FAILURES FOUND"
            should_complete = False
        else:
            status_text     = f"{warnings} WARNINGS"
            should_complete = False

        date_str     = datetime.now().strftime('%Y-%m-%d')
        today_date   = datetime.now().strftime('%Y-%m-%d')
        ticket_title = f"[AI Test Agent] Daily Test Report - {date_str} - {status_text}"

        description_text = f"""AI Test Agent - Daily Automated Test Report
Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Agent: {agent.agent_name}
Version: {agent.version}
Website: https://www.automationanywhere.com/

TEST SUMMARY
Total Tests:  {total}
Passed:       {passed}
Failed:       {failed}
Warnings:     {warnings}
Pass Rate:    {pass_rate:.1f}%
Issues Found: {issues_count}

TEST RESULTS
"""
        for i, result in enumerate(agent.test_results, 1):
            status_icon = (
                "[PASS]" if result['status'] == 'PASS' else
                "[FAIL]" if result['status'] == 'FAIL' else
                "[WARN]"
            )
            description_text += f"{i}. {result['test_name']} - {status_icon} - {result['message']}\n"

        if agent.issues_found:
            description_text += "\nISSUES FOUND\n"
            for i, issue in enumerate(agent.issues_found, 1):
                description_text += f"\nIssue #{i}: {issue['test_name']}\n"
                description_text += f"  - Type: {issue['issue_type']}\n"
                description_text += f"  - Description: {issue['description']}\n"

        description_text += (
            "\n\nThis ticket was automatically created by AI Test Agent v2.0\n"
            "NVISH Solutions | AAI-Web Team"
        )

        url  = f"{JIRA_CONFIG['base_url']}/rest/api/3/issue"
        auth = HTTPBasicAuth(JIRA_CONFIG['email'], JIRA_CONFIG['api_token'])
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }

        active_sprint_id = get_active_sprint_id("AWT", auth)

        payload = {
            "fields": {
                "project":    {"key": "AWT"},
                "summary":    ticket_title,
                "description": {
                    "type": "doc",
                    "version": 1,
                    "content": [
                        {
                            "type": "codeBlock",
                            "attrs": {"language": "text"},
                            "content": [{"type": "text", "text": description_text}]
                        }
                    ]
                },
                "issuetype":  {"name": "Task"},
                "priority":   {"name": "High"},
                "assignee":   {"accountId": "5ea6d2944e41b90b7bce0c0c"},
                "customfield_15267": {"value": "P1 - 1 to 3 days - Assets"},
                "customfield_15268": {"value": "Yes"},
                "customfield_15274": today_date,
                "timetracking": {
                    "originalEstimate":  "1h",
                    "remainingEstimate": "5m"
                }
            }
        }

        if JIRA_CONFIG.get('labels'):
            payload['fields']['labels'] = JIRA_CONFIG['labels']

        print(f"   Project:                 AWT (AAI-Web Team)")
        print(f"   Work Type:               Task")
        print(f"   Priority:                High")
        print(f"   Assignee:                Ritin Verma")
        print(f"   Sprint:                  {'Will set after creation: #' + str(active_sprint_id) if active_sprint_id else 'No active sprint found'}")
        print(f"   Original Estimate:       1h")
        print(f"   Remaining Estimate:      5m")
        print(f"   Digital Marketing SLA:   P1 - 1 to 3 days - Assets")
        print(f"   Product Backlog Ready:   Yes")
        print(f"   Estimated Delivery Date: {today_date}")
        print(f"   Creating ticket...")

        response = requests.post(url, json=payload, headers=headers, auth=auth)

        if response.status_code == 201:
            ticket_data = response.json()
            ticket_key  = ticket_data['key']
            ticket_url  = f"{JIRA_CONFIG['base_url']}/browse/{ticket_key}"

            print(f"\n   Ticket created successfully!")
            print(f"   Ticket: {ticket_key}")
            print(f"   URL:    {ticket_url}")

            sprint_set = False
            if active_sprint_id:
                sprint_set = set_sprint_on_jira_ticket(ticket_key, active_sprint_id, auth)
                if not sprint_set:
                    print(f"   WARNING: Sprint could not be confirmed.")
                else:
                    time.sleep(2)

            attach_files_to_jira(ticket_key, reports, auth)

            if should_complete:
                print(f"\n   All tests passed - transitioning to COMPLETED...")
                success = transition_jira_ticket(ticket_key, "Completed", auth)
                final_status = "Completed" if success else "Backlog"
            else:
                print(f"\n   Issues found - ticket remains in current status")
                final_status = "In Progress"

            return {
                "ticket_key": ticket_key,
                "ticket_url": ticket_url,
                "status":     final_status
            }
        else:
            print(f"\n   Failed to create ticket: {response.status_code}")
            print(f"   Response: {response.text[:500]}")
            try:
                error_data = response.json()
                if 'errors' in error_data:
                    print(f"   Errors: {error_data['errors']}")
            except Exception:
                pass
            return None

    except Exception as e:
        print(f"\n   Jira error: {str(e)}")
        traceback.print_exc()
        return None


# =============================================================================
# EMAIL FUNCTIONS
# =============================================================================
def send_email_report(reports, agent):
    print("\n" + "=" * 70)
    print("SENDING EMAIL REPORT")
    print("=" * 70)

    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To']   = ", ".join(EMAIL_CONFIG['recipients'])

        summary = reports['summary']
        if summary['failed'] > 0:
            subject = f"[ACTION REQUIRED] AI Test Agent Report - {summary['failed']} Failures - {datetime.now().strftime('%Y-%m-%d')}"
        elif summary['warnings'] > 0:
            subject = f"[WARNING] AI Test Agent Report - {summary['warnings']} Warnings - {datetime.now().strftime('%Y-%m-%d')}"
        else:
            subject = f"[SUCCESS] AI Test Agent Report - All Tests Passed - {datetime.now().strftime('%Y-%m-%d')}"

        msg['Subject'] = subject

        body = f"""
        <html>
        <body style="font-family: Arial, sans-serif;">
            <h1 style="color: #667eea;">AI Test Agent - Daily Report</h1>
            <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <p><strong>Agent:</strong> {agent.agent_name}</p>
            <h2>Test Summary</h2>
            <ul>
                <li>Total: {summary['total']}</li>
                <li style="color: green;">Passed: {summary['passed']}</li>
                <li style="color: red;">Failed: {summary['failed']}</li>
                <li style="color: orange;">Warnings: {summary['warnings']}</li>
                <li><strong>Pass Rate: {summary['pass_rate']:.1f}%</strong></li>
            </ul>
            <p>Please see attached reports for details.</p>
            <p><em>Generated by AI Test Agent v2.0 - NVISH Solutions</em></p>
        </body>
        </html>
        """

        msg.attach(MIMEText(body, 'html'))

        if reports.get('html_report') and os.path.exists(reports['html_report']):
            with open(reports['html_report'], 'rb') as f:
                attachment = MIMEBase('application', 'octet-stream')
                attachment.set_payload(f.read())
                encoders.encode_base64(attachment)
                attachment.add_header('Content-Disposition',
                                      f'attachment; filename=test_report_{datetime.now().strftime("%Y%m%d")}.html')
                msg.attach(attachment)
            print(f"   Attached: HTML Report")

        if reports.get('issue_document') and os.path.exists(reports['issue_document']):
            with open(reports['issue_document'], 'rb') as f:
                attachment = MIMEBase('application', 'octet-stream')
                attachment.set_payload(f.read())
                encoders.encode_base64(attachment)
                attachment.add_header('Content-Disposition',
                                      f'attachment; filename=issue_report_{datetime.now().strftime("%Y%m%d")}.docx')
                msg.attach(attachment)
            print(f"   Attached: Issue Document")

        print(f"\n   Connecting to SMTP server...")
        server = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        server.starttls()
        server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])

        print(f"   Sending email to {len(EMAIL_CONFIG['recipients'])} recipients...")
        server.send_message(msg)
        server.quit()

        print(f"\n   EMAIL SENT SUCCESSFULLY!")
        print(f"   Recipients: {', '.join(EMAIL_CONFIG['recipients'])}")
        return True

    except Exception as e:
        print(f"\n   EMAIL FAILED: {str(e)}")
        traceback.print_exc()
        return False


# =============================================================================
# MAIN EXECUTION FUNCTIONS
# =============================================================================
def run_scheduled_test():
    print("\n" + "=" * 70)
    print(f"SCHEDULED TEST EXECUTION - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70)

    try:
        agent   = AITestAgentScheduled(agent_name="NVISH-QA-Agent-001")
        reports = agent.run_test_suite(mode="all")

        summary = reports.get("summary", {})
        print("\nSCHEDULED RUN SUMMARY:")
        print(f"   Total  : {summary.get('total', 0)}")
        print(f"   Passed : {summary.get('passed', 0)}")
        print(f"   Failed : {summary.get('failed', 0)}")
        print(f"   Warn   : {summary.get('warnings', 0)}")

        send_email_report(reports, agent)
        jira_result = create_jira_ticket(reports, agent)

        if jira_result:
            print("\n" + "=" * 70)
            print("JIRA TICKET SUMMARY")
            print("=" * 70)
            print(f"   Ticket: {jira_result['ticket_key']}")
            print(f"   Status: {jira_result['status']}")
            print(f"   URL:    {jira_result['ticket_url']}")
            print("=" * 70)

        try:
            import webbrowser
            if reports.get('html_report'):
                webbrowser.open(reports['html_report'])
        except:
            pass

        print("\n" + "=" * 70)
        print("EXECUTION COMPLETE")
        print("=" * 70)

    except Exception as e:
        print(f"ERROR: {str(e)}")
        traceback.print_exc()


def run_once():
    run_scheduled_test()


def run_scheduled():
    # Make sure Jira is enabled for scheduled runs
    JIRA_CONFIG["enabled"] = True

    # Clear any leftover jobs from previous runs
    schedule.clear()

    # Register the daily job
    schedule.every().day.at(SCHEDULE_TIME).do(run_scheduled_test)

    # Calculate time until next run
    next_run = schedule.next_run()
    now      = datetime.now()
    delta    = next_run - now
    hours, remainder = divmod(int(delta.total_seconds()), 3600)
    minutes          = remainder // 60

    print("\n" + "=" * 70)
    print("AI TEST AGENT - SCHEDULER STARTED")
    print("=" * 70)
    print(f"   Scheduled Time : {SCHEDULE_TIME} IST (daily)")
    print(f"   Current Time   : {now.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"   Next Run At    : {next_run.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"   Time Until Run : {hours}h {minutes}m")
    print(f"   Jira Enabled   : True")
    print(f"   Project        : {JIRA_CONFIG.get('project_key', 'N/A')} (AAI-Web Team)")
    print("=" * 70)
    print("\nNOTE: Scheduler uses your SYSTEM clock.")
    print(f"      If system clock is UTC, change SCHEDULE_TIME to '13:30'.")
    print("\nPress Ctrl+C to stop.\n")

    # Tick every second so the job fires on time (not up to 60s late)
    while True:
        schedule.run_pending()
        time.sleep(1)


# =============================================================================
# MAIN ENTRY POINT — called by Streamlit UI via run_tests()
# FIX: now sends email and creates Jira ticket when requested
# =============================================================================
def run_tests(mode: str = "all", with_jira: bool = False, run_headless: bool = True, send_email: bool = True):
    # ------------------------------------------------------------------ #
    # BUG FIX: previously this function ONLY ran the test suite and       #
    # returned results. Email and Jira were never triggered from the UI.  #
    # ------------------------------------------------------------------ #
    JIRA_CONFIG["enabled"] = with_jira
    agent   = AITestAgentScheduled(agent_name="QA-Agent-UI", run_headless=run_headless)
    reports = agent.run_test_suite(mode=mode)

    # Send email only if enabled
    if send_email:
        print("\n--- Sending email report ---")
        send_email_report(reports, agent)
    else:
        print("\n--- Email report skipped ---")

    # Create Jira ticket only when requested via UI toggle
    jira_result = None
    if with_jira:
        print("\n--- Creating Jira ticket ---")
        jira_result = create_jira_ticket(reports, agent)

    # Attach jira result to reports dict so Streamlit can display it
    reports["jira_result"] = jira_result

    return reports


if __name__ == "__main__":
    import sys
    import webbrowser

    # ------------------------------------------------------------------
    # Usage:
    #   python ai_test_agent.py                  → start scheduler (waits for 19:00 IST)
    #   python ai_test_agent.py run-now          → run tests immediately once
    #   python ai_test_agent.py run-now with-jira → run now + create Jira ticket
    # ------------------------------------------------------------------

    arg1 = sys.argv[1].lower() if len(sys.argv) > 1 else "schedule"
    arg2 = sys.argv[2].lower() if len(sys.argv) > 2 else "no-jira"

    if arg1 == "run-now":
        # Immediate one-off run (used for manual testing / Streamlit)
        with_jira = (arg2 == "with-jira")
        print(f"\nRunning tests immediately (with_jira={with_jira})...")
        results = run_tests(mode="all", with_jira=with_jira)

        html_report = results.get("html_report")
        if html_report and os.path.exists(html_report):
            try:
                webbrowser.open_new_tab(html_report)
                print(f"\nOpened report: {html_report}")
            except Exception as e:
                print(f"\nCould not open report automatically: {str(e)[:80]}")
    else:
        # Default: start the scheduler — waits for 19:00 IST daily
        run_scheduled()
